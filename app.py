import streamlit as st
import os
import re
import pandas as pd
from datetime import datetime
from google.cloud import vision
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from google.oauth2 import service_account

st.set_page_config(page_title="Legal_AI: 문서 자동화", layout="wide")

st.markdown("""
    <style>
    .info-box { background-color: #ffffcc; padding: 20px; border-radius: 12px; box-shadow: 2px 2px 10px rgba(0,0,0,0.1); border-left: 5px solid #ff4b4b; margin-bottom: 20px; }
    .main-title { color: #2c3e50; font-weight: bold; }
    </style>
    """, unsafe_allow_html=True)

with st.sidebar:
    st.header("🏢 지배구조 설정")
    company_name_input = st.text_input("수동 회사명 입력 (AI가 인식 못할 때만 사용)", "")
    st.divider()
    is_listed = st.checkbox("📈 상장회사(Listed) 적용", value=False)
    term_setting = st.radio("일반 이사 임기 (정관 기준)", [2, 3], index=1)

st.markdown("<h1 class='main-title'>⚖️ Legal_AI: 문서 자동 생성 시스템</h1>", unsafe_allow_html=True)

def create_minutes_docx(company_name, name, title, action_type, deadline_date):
    doc = Document()
    
    heading = doc.add_heading('임 시 주 주 총 회 의 사 록', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"1. 일  시 : {deadline_date}  오전 10시 00분")
    doc.add_paragraph(f"2. 장  소 : {company_name} 본점 회의실")
    
    doc.add_paragraph(f"3. 발행주식의 총수 :               주 (의결권 있는 주식수 :               주)")
    doc.add_paragraph(f"4. 출석주주 수 :               명")
    doc.add_paragraph(f"5. 출석주주의 주식수 :               주")
    
    doc.add_paragraph("\n대표이사 OOO은 의장석에 등단하여 위와 같이 법정수에 달하는 주주가 출석하였으므로 본 총회가 적법하게 성립되었음을 알리고 개회를 선언하다. 이어 다음 의안을 부의하고 심의를 구하다.\n")
    
    doc.add_heading(f'제 1 호 의안 : {title} 선임의 건', level=1)
    
    if action_type == "교체":
        doc.add_paragraph(
            f"의장은 상법 및 당사 정관에 의거하여 기존 {title}의 임기(또는 상법상 최대 재임기간)가 만료됨에 따라 "
            f"신임 {title}를 선임하여야 할 필요성을 설명하고, 그 후보자로 {name}을(를) 추천한 바, "
            f"출석한 주주 전원이 이의 없이 찬성하여 다음과 같이 선임하다."
        )
    else:
        doc.add_paragraph(
            f"의장은 당사 {title} {name}의 임기가 만료됨에 따라, 회사의 지속적인 업무 수행을 위하여 "
            f"재선임(중임)할 필요가 있음을 설명하고 그 가부를 물은 바, "
            f"출석한 주주 전원이 이의 없이 찬성하여 다음과 같이 중임(선임)하다."
        )
    
    p_result = doc.add_paragraph()
    p_result.add_run(f"- 선임된 {title} : {name}\n").bold = True
    p_result.add_run(f"- 찬성 주식수 :               주 (출석 의결권의 100%)")
    
    doc.add_paragraph(f"\n위 선임된 {name}은(는) 즉석에서 그 취임을 승낙하다.")
    
    doc.add_paragraph("\n의장이 이상으로 총회의 목적사항을 전부 종료하였으므로 폐회를 선언하다.")
    doc.add_paragraph("(폐회시간  11시 00분)")
    
    doc.add_paragraph(f"\n위 결의를 명확히 하기 위하여 이 의사록을 작성하고 의장과 출석한 이사가 기명날인하다.\n")
    
    doc.add_paragraph(f"202X년  XX월  XX일")
    doc.add_paragraph(f"회사명 : {company_name}")
    
    doc.add_paragraph("의장  대표이사          O O O  (인)")
    doc.add_paragraph("      사내이사          O O O  (인)")
    doc.add_paragraph("      사내이사          O O O  (인)")
    doc.add_paragraph("      감    사          O O O  (인)")
    
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

def get_actionable_dates_enterprise(first_appt_str, latest_appt_str, title, term_years, is_listed):
    try:
        nums_first = re.findall(r'\d+', first_appt_str)
        nums_latest = re.findall(r'\d+', latest_appt_str)
        first_date = datetime(int(nums_first[0]), int(nums_first[1]), int(nums_first[2]))
        latest_date = datetime(int(nums_latest[0]), int(nums_latest[1]), int(nums_latest[2]))
        
        base_expiry = latest_date.replace(year=latest_date.year + term_years)
        expiry_str = base_expiry.strftime('%Y-%m-%d')
        action_type = "중임/선임"
        
        if is_listed and "사외이사" in title:
            max_limit_date = first_date.replace(year=first_date.year + 6)
            if base_expiry >= max_limit_date:
                action_type = "교체"
                return f"{max_limit_date.strftime('%Y-%m-%d')}", f"🚨 무조건 교체", action_type

        if base_expiry.month <= 3:
            return f"{expiry_str} (정기주총 연장)", f"{base_expiry.year}년 3월 정기주주총회 종결 시", action_type
        else:
            return f"{expiry_str}", f"{expiry_str} 이전 (임시주총 요망)", action_type
    except Exception as e:
        return None, None, "오류"

uploaded_file = st.file_uploader("법인등기부 이미지를 업로드하세요", type=["png", "jpg", "jpeg"])

if uploaded_file is not None:
    with st.spinner('AI가 지배구조 스캔 및 회사명을 추출 중입니다...'):
        # [핵심 변경] 클라우드 보안 인증 로직 적용
        try:
            key_dict = json.loads(st.secrets["google_key"])
            credentials = service_account.Credentials.from_service_account_info(key_dict)
            client = vision.ImageAnnotatorClient(credentials=credentials)
        except Exception as e:
            st.error("보안 키가 설정되지 않았습니다. Streamlit 클라우드 설정(Secrets)을 확인해주세요.")
            st.stop()

        content = uploaded_file.getvalue()
        image = vision.Image(content=content)
        response = client.document_text_detection(image=image)
        text = response.full_text_annotation.text

        company_match = re.search(r"상\s*호\s*([가-힣A-Za-z\(\)\s]+?)(?=\n|본|지)", text)
        detected_company = company_match.group(1).strip() if company_match else "회사명 인식불가"
        final_company_name = company_name_input if company_name_input else detected_company

        st.success(f"분석 완료! (AI 인식 회사명: {final_company_name})")
        col1, col2 = st.columns([1, 1.2])
        
        with col1:
            st.image(uploaded_file, use_container_width=True)
        
        with col2:
            parts = re.split(r"(사내이사|대표이사|감사위원|감사|사외이사|기타비상무이사)", text)
            blocks = [(parts[i], parts[i+1]) for i in range(1, len(parts)-1, 2)]
            
            registry_db = {}
            for title, content_text in blocks:
                name_match = re.search(r"^\s*([가-힣A-Za-z\(\)\s,\-]+?)(?:\d{6}-|\d{4}\s*년\s*\d{1,2}\s*월)", content_text)
                raw_name = name_match.group(1) if name_match else content_text.split('\n')[0][:20]
                name = re.sub(r"^[-*.\s]+", "", raw_name).strip() 
                
                id_match = re.search(r"(\d{6}-[\d\*]{7}|\d{4}\s*년\s*\d{1,2}\s*월\s*\d{1,2}\s*일)", content_text)
                unique_id = id_match.group(1).replace(" ", "") if id_match else re.sub(r'[^가-힣A-Za-z]', '', name)[-5:]
                unique_key = f"{title}_{unique_id}"
                
                term_start_dates = re.findall(r"(\d{4}\s*년\s*\d{1,2}\s*월\s*\d{1,2}\s*일)\s*(?:취임|중임)", content_text)
                resign_dates = re.findall(r"(\d{4}\s*년\s*\d{1,2}\s*월\s*\d{1,2}\s*일)\s*(?:사임|퇴임)", content_text)
                
                if unique_key not in registry_db:
                    registry_db[unique_key] = {'status': 'unknown', 'first_appt': None}
                
                registry_db[unique_key]['title'] = title
                registry_db[unique_key]['name'] = name
                
                if term_start_dates:
                    if not registry_db[unique_key]['first_appt']:
                        registry_db[unique_key]['first_appt'] = term_start_dates[0]
                    registry_db[unique_key]['latest_appt'] = term_start_dates[-1]
                
                if resign_dates:
                    registry_db[unique_key]['status'] = 'resigned'
                else:
                    if registry_db[unique_key].get('status') != 'resigned':
                        registry_db[unique_key]['status'] = 'active'
            
            for key, data in registry_db.items():
                if data.get('status') == 'active' and data.get('latest_appt'):
                    title = data['title']
                    name = data['name']
                    
                    combined_expiry, action_deadline, action_type = get_actionable_dates_enterprise(
                        data['first_appt'], data['latest_appt'], title, term_setting, is_listed
                    )
                    
                    with st.container():
                        col_info, col_btn = st.columns([3, 1])
                        with col_info:
                            st.markdown(f'<div class="info-box"><span class="hover-red" style="font-size:1.2em; font-weight:bold;">👤 {title} {name}</span><br><small style="color:gray;">기산일: {data["latest_appt"]}</small><br><div style="margin-top:10px;">📅 <b>임기 만료:</b> {combined_expiry}<br>🚀 <b>선임 기한: <span style="color:#d32f2f;">{action_deadline}</span></b></div></div>', unsafe_allow_html=True)
                        with col_btn:
                            doc_bytes = create_minutes_docx(final_company_name, name, title, action_type, action_deadline)
                            safe_company_name = re.sub(r'[^가-힣A-Za-z0-9]', '_', final_company_name)
                            safe_file_name = f"{safe_company_name}_{name}_의사록.docx"
                            
                            st.download_button(
                                label=f"📄 의사록(Word) 생성",
                                data=doc_bytes,
                                file_name=safe_file_name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )